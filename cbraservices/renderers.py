from rest_framework import renderers
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from io import BytesIO
from localflavor.us import us_states


class DOCXRenderer(renderers.BaseRenderer):
    media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    format = 'docx'
    render_style = 'binary'
    document = None #This should be a docx object

    def render(self, data, media_type=None, renderer_context=None):
        return self.document


class FinalLetterDOCXRenderer(DOCXRenderer):

    def render(self, data, media_type=None, renderer_context=None):
        DOC_FONT_TYPE = 'Times New Roman'
        DOC_FONT_SMALL = Pt(8)
        DOC_FONT_LARGE = Pt(12)
        DOC_LINE_SPACING = 1

        # case fields
        id = str(data[0]['id'])
        case_reference = data[0]['case_reference'] or ""
        request_date = data[0]['request_date'] or ""
        if request_date:
            request_date = datetime.strptime(data[0]['request_date'], '%Y-%m-%d').strftime('%B %-d, %Y')
        determination = data[0]['determination'] or ""
        determination_string = data[0]['determination_string'] or ""
        cbrs_unit = data[0]['cbrs_unit_string'] or ""
        system_unit_type = data[0]['system_unit_type'] or ""
        prohibition_date = data[0]['prohibition_date'] or ""
        if prohibition_date:
            prohibition_date = datetime.strptime(data[0]['prohibition_date'], '%Y-%m-%d').strftime('%B %-d, %Y')
        map_number = str(data[0]['map_number']) or ""
        cbrs_map_date = data[0]['cbrs_map_date'] or ""
        if cbrs_map_date:
            cbrs_map_date = datetime.strptime(data[0]['cbrs_map_date'], '%Y-%m-%d').strftime('%B %-d, %Y')
        final_letter_recipient = data[0]['final_letter_recipient'] or ""

        # property fields
        property_legal_description = data[0]['legal_description'] or ""
        property_subdivision = data[0]['subdivision'] or ""
        property_policy_number = data[0]['policy_number'] or ""
        property_street = data[0]['property_street'] or ""
        property_unit = data[0]['property_unit'] or ""
        property_city = data[0]['property_city'] or ""
        property_state = data[0]['property_state'] or ""
        property_zipcode = data[0]['property_zipcode'] or ""

        if property_state != "":
            property_state = str(next(name for abbrev, name in us_states.US_STATES if abbrev == property_state))

        # requester fields
        requester_salutation = data[0]['salutation'] or ""
        requester_first_name = data[0]['first_name'] or ""
        requester_last_name = data[0]['last_name'] or ""
        requester_street = data[0]['requester_street'] or ""
        requester_unit = data[0]['requester_unit'] or ""
        requester_city = data[0]['requester_city'] or ""
        requester_state = data[0]['requester_state'] or ""
        requester_zipcode = data[0]['requester_zipcode'] or ""

        if requester_state != "":
            requester_state = str(next(name for abbrev, name in us_states.US_STATES if abbrev == requester_state))

        # letter content sections

        referal = "In Reply Refer To\nFWS/DBTS-BGMTS"

        requester_full_address = "\n" + requester_salutation + " " + requester_first_name + " " + requester_last_name + "\n"
        if requester_unit != "":
            requester_full_address += requester_unit + " "
        requester_full_address += requester_street + "\n"
        requester_full_address += requester_city + ", " + requester_state + " " + requester_zipcode + "\n\n"

        salutation = "Dear " + requester_salutation + " " + requester_last_name + ","

        intro = "The U.S. Fish and Wildlife Service (Service) has reviewed the request dated " + request_date + ","
        intro += " Case " + case_reference + ", for a determination as to whether the following property"
        intro += " is within a System Unit or an Otherwise Protected Area (OPA) of the John H. Chafee"
        intro += " Coastal Barrier Resources System (CBRS). "
        if property_policy_number != "":
            intro += "The flood insurance policy number for the submitted request is " + property_policy_number + "."

        property_address = "Address:\t\t\t"
        if property_unit != "":
            property_address += property_unit + " "
        property_address += property_street + "\n\t\t\t\t" + property_city + ", " + property_state + " " + property_zipcode

        legal_description = "Legal Description:\t\t"
        if property_legal_description != "":
            legal_description += property_legal_description + " "
        elif property_subdivision != "":
            legal_description += property_subdivision
        else:
            legal_description += "(none submitted)"

        details = "We compared the location of the property above, as depicted on the information that was provided,"
        details += " to the official CBRS map for the area, numbered " + map_number + ", dated " + cbrs_map_date + ". "

        # 1:In, 2:Out, 3:Partially In; Structure In, 4:Partially In; Structure Out, 5:Partially In/No Structure
        if determination == 1:
            details += "This property is within " + system_unit_type + " Unit " + cbrs_unit + "."
            details += " The prohibition date on Federal flood insurance is " + prohibition_date + "."
        elif determination == 2:
            details += "This property is not located within a System Unit or an OPA of the CBRS."
        elif determination == 3:
            details += "This property is partially located within " + system_unit_type + " Unit " + cbrs_unit + "."
            details += " The existing structure located on the property is within Unit " + cbrs_unit + "."
            details += " The prohibition date on Federal flood insurance is " + prohibition_date + "."
        elif determination == 4:
            details += "This property is partially located within " + system_unit_type + " Unit " + cbrs_unit + "."
            details += " The existing structure located on the property is not within Unit " + cbrs_unit + "."
        elif determination == 5:
            details += "This property is partially located within " + system_unit_type + " Unit " + cbrs_unit + "."
            details += " There is no existing structure located on the property. The prohibition date on Federal "
            details += "flood insurance for the portion of the property within the CBRS is " + prohibition_date + "."
        else:
            details += "A determination has not been made."

        closing = "We hope this information is helpful. Additional information concerning the CBRS can be found"
        closing += " on our website at https://www.fws.gov/cbra/. If you have any additional questions, please contact"
        closing += " Ms. Terri Fish, Program Specialist, at (703) 358-2171.\n"

        signature = "\t\t\t\t\t\tSincerely,\n\n\n\n\n"
        signature += "\t\t\t\t\t\tJonathan Phinney, PhD\n"
        signature += "\t\t\t\t\t\tChief, Branch of Geospatial Mapping and\n"
        signature += "\t\t\t\t\t\tTechnical Support\n"

        cc = "cc:\t"
        if final_letter_recipient != "":
            cc += final_letter_recipient

        # assemble the content sections into a document with proper formatting

        document = Document()
        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.line_spacing = DOC_LINE_SPACING
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1.6)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

        p1 = document.add_paragraph().add_run(referal)
        p1.font.name = DOC_FONT_TYPE
        p1.font.size = DOC_FONT_SMALL

        p2 = document.add_paragraph().add_run(requester_full_address)
        p2.font.name = DOC_FONT_TYPE
        p2.font.size = DOC_FONT_LARGE

        p3 = document.add_paragraph().add_run(salutation)
        p3.font.name = DOC_FONT_TYPE
        p3.font.size = DOC_FONT_LARGE

        p4 = document.add_paragraph().add_run(intro)
        p4.font.name = DOC_FONT_TYPE
        p4.font.size = DOC_FONT_LARGE

        p5 = document.add_paragraph().add_run(property_address)
        p5.font.name = DOC_FONT_TYPE
        p5.font.size = DOC_FONT_LARGE

        p6 = document.add_paragraph().add_run(legal_description)
        p6.font.name = DOC_FONT_TYPE
        p6.font.size = DOC_FONT_LARGE

        p7 = document.add_paragraph().add_run(details)
        p7.font.name = DOC_FONT_TYPE
        p7.font.size = DOC_FONT_LARGE

        p8 = document.add_paragraph().add_run(closing)
        p8.font.name = DOC_FONT_TYPE
        p8.font.size = DOC_FONT_LARGE

        p9 = document.add_paragraph().add_run(signature)
        p9.font.name = DOC_FONT_TYPE
        p9.font.size = DOC_FONT_LARGE

        p10 = document.add_paragraph().add_run(cc)
        p10.font.name = DOC_FONT_TYPE
        p10.font.size = DOC_FONT_LARGE

        #document.add_paragraph()

        docx_buffer = BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)
        self.document = docx_buffer

        return super(FinalLetterDOCXRenderer, self).render(data, media_type, renderer_context)
