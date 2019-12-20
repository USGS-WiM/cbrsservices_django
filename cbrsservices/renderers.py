from rest_framework import renderers
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from io import BytesIO
from localflavor.us import us_states
from django.conf import settings
from rest_framework_csv.renderers import CSVRenderer
import re
from docx.shared import RGBColor


class PaginatedCSVRenderer (CSVRenderer):
    results_field = 'results'

    def render(self, data, *args, **kwargs):
        if not isinstance(data, list):
            data = data.get(self.results_field, [])
        return super(PaginatedCSVRenderer, self).render(data, *args, **kwargs)

class WorkbenchCSVRenderer (CSVRenderer):
    header = ['id', 'case_reference', 'status', 'prohibition_date', 'cbrs_unit_string', 'request_date', 'final_letter_date', 'determination_string',
        'property_string', 'tags', 'duplicate', 'distance', 'analyst_string', 'analyst_signoff_date', 'qc_reviewer_string',
        'qc_reviewer_signoff_date', 'final_letter_date', 'requester_string', 'requester_organization', 'requester_email', 'requester_address',
        'priority', 'on_hold', 'invalid', 'hard_copy_map_reviewed']
    labels = {
        'id': 'Case ID',
        'case_reference': 'Case Reference',
        'status': 'Status',
        'prohibition_date': 'Prohibition Date',
        'cbrs_unit_string': 'CBRS Unit',
        'request_date': 'Request Date',
        'final_letter_date': 'Final Letter Date',
        'determination_string': 'Determination',
        'property_string': 'Property',
        'tags': 'Tags',
        'duplicate': 'Duplicate',
        'distance': 'Distance',
        'analyst_string': 'Analyst',
        'analyst_signoff_date': 'Analyst Signoff Date',
        'qc_reviewer_string': 'QC Reviewer',
        'qc_reviewer_signoff_date': 'QC Reviewer Signoff Date',
        'final_letter_date': 'Final Letter Date',
        'requester_string': 'Requester',
        'requester_organization': 'Requester Organization',
        'requester_email': 'Requester Email',
        'requester_address': 'Requester Address',
        'priority': 'Priority',
        'on_hold': 'On Hold',
        'invalid': 'Invalid',
        'hard_copy_map_reviewed': 'Hard Copy Map Reviewed',
    }

class ReportCasesByUnitCSVRenderer (PaginatedCSVRenderer):
    header = ['id', 'status', 'prohibition_date', 'cbrs_unit_string', 'request_date', 'final_letter_date',
              'determination_string', 'street_address', 'tags', 'comments', 'case_number', 'case_reference',
              'duplicate', 'property_string', 'map_number_string', 'cbrs_map_date', 'distance', 'fws_fo_received_date',
              'fws_hq_received_date', 'close_date', 'final_letter_recipient', 'analyst_string', 'analyst_signoff_date',
              'qc_reviewer_string', 'qc_reviewer_signoff_date', 'priority', 'on_hold', 'invalid', 'casefiles',
              'created_by_string', 'modified_by_string', 'hard_copy_map_reviewed']
    labels = {
        'id': 'Case ID',
        'status': 'Status',
        'prohibition_date': 'Prohibition Date',
        'cbrs_unit_string': 'CBRS Unit',
        'request_date': 'Request Date',
        'final_letter_date': 'Final Letter Date',
        'determination_string': 'Determination',
        'street_address': 'Street Address',
        'tags': 'Tags',
        'comments': 'Comments',
        'case_number': 'Case Number',
        'case_reference': 'Case Reference',
        'duplicate': 'Duplicate',
        'property_string': 'Property',
        'map_number_string': 'Map Number',
        'cbrs_map_date': 'CBRS Map Date',
        'distance': 'Distance',
        'fws_fo_received_date': 'FWS FO Received Date',
        'fws_hq_received_date': 'FWS QH Received Date',
        'close_date': 'Close Date',
        'final_letter_recipient': 'Final Letter Recipient',
        'analyst_string': 'Analyst',
        'analyst_signoff_date': 'Analyst Signoff Date',
        'qc_reviewer_string': 'QC Reviewer',
        'qc_reviewer_signoff_date': 'QC Reviewer Signoff Date',
        'priority': 'Priority',
        'on_hold': 'On Hold',
        'invalid': 'Invalid',
        'casefiles': 'Casefiles',
        'created_by_string': 'Created By',
        'modified_by_string': 'Modified By',
        'hard_copy_map_reviewed': 'Hard Copy Map Reviewed'
    }


class ReportCasesForUserCSVRenderer (PaginatedCSVRenderer):
    header = ['id', 'status', 'prohibition_date', 'cbrs_unit_string', 'request_date', 'final_letter_date',
              'determination_string', 'street_address', 'tags', 'comments', 'case_number', 'case_reference',
              'duplicate', 'property_string', 'map_number_string', 'cbrs_map_date', 'distance', 'fws_fo_received_date',
              'fws_hq_received_date', 'close_date', 'final_letter_recipient', 'analyst_string', 'analyst_signoff_date',
              'qc_reviewer_string', 'qc_reviewer_signoff_date', 'priority', 'on_hold', 'invalid', 'casefiles',
              'created_by_string', 'modified_by_string', 'hard_copy_map_reviewed']
    labels = {
        'id': 'Case ID',
        'status': 'Status',
        'prohibition_date': 'Prohibition Date',
        'cbrs_unit_string': 'CBRS Unit',
        'request_date': 'Request Date',
        'final_letter_date': 'Final Letter Date',
        'determination_string': 'Determination',
        'street_address': 'Street Address',
        'tags': 'Tags',
        'comments': 'Comments',
        'case_number': 'Case Number',
        'case_reference': 'Case Reference',
        'duplicate': 'Duplicate',
        'property_string': 'Property',
        'map_number_string': 'Map Number',
        'cbrs_map_date': 'CBRS Map Date',
        'distance': 'Distance',
        'fws_fo_received_date': 'FWS FO Received Date',
        'fws_hq_received_date': 'FWS QH Received Date',
        'close_date': 'Close Date',
        'final_letter_recipient': 'Final Letter Recipient',
        'analyst_string': 'Analyst',
        'analyst_signoff_date': 'Analyst Signoff Date',
        'qc_reviewer_string': 'QC Reviewer',
        'qc_reviewer_signoff_date': 'QC Reviewer Signoff Date',
        'priority': 'Priority',
        'on_hold': 'On Hold',
        'invalid': 'Invalid',
        'casefiles': 'Casefiles',
        'created_by_string': 'Created By',
        'modified_by_string': 'Modified By',
        'hard_copy_map_reviewed': 'Hard Copy Map Reviewed'
    }


class ReportDaysToResolutionCSVRenderer (PaginatedCSVRenderer):
    header = ['id', 'case_reference', 'request_date', 'close_date', 'close_days']
    labels = {
        'id': 'Case ID',
        'case_reference': 'Case Reference',
        'request_date': 'Request Date',
        'close_date': 'Close Date',
        'close_days': 'Days to Close'
    }


class ReportDaysToEachStatusCSVRenderer (PaginatedCSVRenderer):
    header = ['id', 'case_reference', 'request_date', 'analyst_signoff_date', 'analyst_days',
              'qc_reviewer_signoff_date', 'qc_reviewer_days', 'final_letter_date', 'final_letter_days',
              'close_date', 'close_days']
    labels = {
        'id': 'Case ID',
        'case_reference': 'Case Reference',
        'request_date': 'Request Date',
        'analyst_signoff_date': 'Awaiting QC',
        'analyst_days': 'Days to QC',
        'qc_reviewer_signoff_date': 'Awaiting Final Letter Date',
        'qc_reviewer_days': 'Days to Awaiting Final Letter',
        'final_letter_date': 'Final Letter Date',
        'final_letter_days': 'Days to Final Letter',
        'close_date': 'Close Date',
        'close_days': 'Days to Close'
    }


class ReportCaseCountCSVRenderer (CSVRenderer):
    header = ['count_received', 'count_awaiting_qc', 'count_awaiting_final_letter',
              'count_closed', 'count_closed_no_final_letter']
    labels = {
        'count_received': 'Count Received',
        'count_awaiting_qc': 'Count Awaiting QC',
        'count_awaiting_final_letter': 'Count Awaiting Final Letter',
        'count_closed': 'Count Closed',
        'count_closed_no_final_letter': 'Count Closed with No Final Letter'
    }


class DOCXRenderer(renderers.BaseRenderer):
    media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    format = 'docx'
    render_style = 'binary'
    document = None

    def render(self, data, media_type=None, renderer_context=None):
        return self.document


class FinalLetterDOCXRenderer(DOCXRenderer):

    def render(self, data, media_type=settings.CONTENT_TYPE_DOCX, renderer_context=None):
        DOC_FONT_TYPE = 'Times New Roman'
        DOC_FONT_SMALL = Pt(8)
        DOC_FONT_LARGE = Pt(12)
        DOC_LINE_SPACING = 1

        case = data[0]

        # case fields
        # id = str(case['id'])
        case_reference = case['case_reference'] if 'case_reference' in case else ""
        request_date = case['request_date'] if 'request_date' in case else ""
        if request_date:
            request_date = datetime.strptime(case['request_date'], '%Y-%m-%d').strftime('%B %d, %Y').replace(" 0", " ")
        determination = case['determination'] if 'determination' in case else ""
        # determination_string = case['determination_string'] if 'determination_string' in case else ""
        cbrs_unit = case['cbrs_unit_string'] if 'cbrs_unit_string' in case else ""
        system_unit_type = case['system_unit_type'] if 'system_unit_type' in case else ""
        if system_unit_type == 'CBRS':
            system_unit_type = 'System'
        prohibition_date = case['prohibition_date'] if 'prohibition_date' in case else ""
        if prohibition_date:
            prohibition_date = datetime.strptime(
                case['prohibition_date'], '%Y-%m-%d').strftime('%B %d, %Y').replace(" 0", " ")
        map_number_string = str(case['map_number_string']) if 'map_number_string' in case else ""
        cbrs_map_date = case['cbrs_map_date'] if 'cbrs_map_date' in case else ""
        if cbrs_map_date:
            cbrs_map_date = datetime.strptime(
                case['cbrs_map_date'], '%Y-%m-%d').strftime('%B %d, %Y').replace(" 0", " ")
        final_letter_recipient = case['final_letter_recipient'] if 'final_letter_recipient' in case else ""

        # property fields
        property_legal_description = case['legal_description'] if 'legal_description' in case else ""
        property_subdivision = case['subdivision'] if 'subdivision' in case else ""
        property_policy_number = case['policy_number'] if 'policy_number' in case else ""
        property_street = case['property_street'] if 'property_street' in case else ""
        property_unit = case['property_unit'] if 'property_unit' in case else ""
        property_city = case['property_city'] if 'property_city' in case else ""
        property_state = case['property_state'] if 'property_state' in case else ""
        property_zipcode = case['property_zipcode'] if 'property_zipcode' in case else ""

        if property_state != "":
            property_state = str(next(name for abbrev, name in us_states.US_STATES if abbrev == property_state))

        # requester fields
        requester_salutation = case['salutation'] if 'salutation' in case else ""
        requester_first_name = case['first_name'] if 'first_name' in case else ""
        requester_last_name = case['last_name'] if 'last_name' in case else ""
        requester_organization = case['requester_organization'] if 'requester_organization' in case else ""
        requester_street = case['requester_street'] if 'requester_street' in case else ""
        requester_unit = case['requester_unit'] if 'requester_unit' in case else ""
        requester_city = case['requester_city'] if 'requester_city' in case else ""
        requester_state = case['requester_state'] if 'requester_state' in case else ""
        requester_zipcode = case['requester_zipcode'] if 'requester_zipcode' in case else ""

        if requester_state != "":
            requester_state = str(next(name for abbrev, name in us_states.US_STATES if abbrev == requester_state))

        # letter content sections

        referal = "In Reply Refer To\nFWS/DBTS-BGMTS"

        requester_full_address = "\n" + requester_salutation + " " + requester_first_name + " " + requester_last_name
        requester_full_address += "\n" + requester_organization + "\n"
        requester_full_address += requester_street
        if requester_unit:
            requester_full_address += ", " + requester_unit + "\n"
        else:
            requester_full_address += "\n"
        if requester_city and requester_state and requester_zipcode:
            requester_full_address += requester_city + ", " + requester_state + " " + requester_zipcode + "\n\n"
        elif not requester_city and requester_state and requester_zipcode:
            requester_full_address += requester_state + " " + requester_zipcode + "\n\n"
        elif not requester_city and not requester_state and requester_zipcode:
            requester_full_address += requester_zipcode + "\n\n"
        elif requester_city and not requester_state and requester_zipcode:
            requester_full_address += requester_city + ", " + requester_zipcode + "\n\n"
        elif requester_city and requester_state and not requester_zipcode:
            requester_full_address += requester_city + ", " + requester_state + "\n\n"
        elif requester_city and not requester_state and not requester_zipcode:
            requester_full_address += requester_city + "\n\n"
        elif not requester_city and requester_state and not requester_zipcode:
            requester_full_address += requester_state + "\n\n"

        salutation = "Dear " + requester_salutation + " " + requester_last_name + ","

        intro = "The U.S. Fish and Wildlife Service is responsible for maintaining the John H. Chafee Coastal Barrier "
        intro += "Resources System (CBRS) maps. We have reviewed your request dated " + request_date + ", "
        intro += "Case Reference # " + case_reference + ", for a determination as to whether the following property "
        intro += "is within a System Unit or an Otherwise Protected Area (OPA) of "
        intro += "the John H. Chafee Coastal Barrier Resources System (CBRS)."
        if property_policy_number != "":
            intro += " The flood insurance policy number for this request is " + property_policy_number + "."

        property_address = "Address:\t\t\t"
        if property_unit != "":
            property_address += property_unit + " "
        property_address += property_street + "\n\t\t\t\t"
        property_address += property_city + ", " + property_state + " " + property_zipcode

        legal_description = ""
        if property_legal_description != "":
            legal_description = "Legal Description:\t\t" + property_legal_description + " "
            if property_subdivision != "":
                legal_description += property_subdivision
        elif property_subdivision != "":
            legal_description = "Legal Description:\t\t" + property_subdivision

        details = "We compared the location of the property above, as depicted on the information that was provided,"
        details += " to the official CBRS map for the area, numbered " + map_number_string
        details += ", dated " + cbrs_map_date + ". "

        bold = ""

        # 1:In, 2:Out, 3:Partially In; Structure In, 4:Partially In; Structure Out, 5:Partially In/No Structure
        su_info = "The Coastal Barrier Resources Act (Pub. L. 97-348) and subsequent amendments "
        su_info += "(16 U.S.C. § 3501 et seq.) prohibit most Federal funding and financial assistance "
        opa_info1 = "within System Units, including flood insurance. The Coastal Barrier Improvement Act "
        opa_info2 = "(Pub. L. 101-591; 42 U.S.C. § 4028) prohibits Federal flood insurance, "
        opa_info2 += "with an exception for structures that are used in a manner consistent with the purpose "
        opa_info2 += "for which the area is protected (e.g., park visitors center, park restroom facilities, etc.). "
        is_opa = cbrs_unit.endswith('P') if cbrs_unit else False
        if determination == 1:
            details += "This property is within " + system_unit_type + " Unit " + cbrs_unit + " of the CBRS. "
            details += opa_info1 + "within OPAs" + opa_info2 if is_opa else su_info + ". "
            details += "\n\nThe prohibition on Federal flood insurance for this property took effect on "
            details += prohibition_date + "."
        elif determination == 2:
            details += "This property is not located within a System Unit or an OPA of the CBRS."
        elif determination == 3:
            details += "This property is partially within " + system_unit_type + " Unit " + cbrs_unit + " of the CBRS. "
            details += "The existing structure on the property is within Unit " + cbrs_unit + ". "
            details += opa_info1 + "for new construction" + opa_info2 if is_opa else su_info + ". "
            details += "\n\nThe prohibition on Federal flood insurance for the portion of this property "
            details += "within the CBRS took effect on " + prohibition_date + "."
        elif determination == 4:
            details += "This property is partially within " + system_unit_type + " Unit " + cbrs_unit
            details += " of the CBRS. Only the portion of the property within Unit "
            if is_opa:
                details += cbrs_unit + " is affected by the Coastal Barrier Improvement Act (CBIA) "
                details += "(Pub. L. 101-591; 42 § U.S.C. 4028). The CBIA prohibits Federal flood insurance "
                details += "within OPAs for new construction, with an exception for structures that are used "
                details += "in a manner consistent with the purpose for which the area is protected "
                details += "(e.g., park visitors center, park restroom facilities, etc.). However, "
                bold = "the existing structure on the property is not within Unit " + cbrs_unit
                bold += " and is therefore not affected by the CBIA restriction on Federal flood insurance."
            else:
                details += cbrs_unit + "is affected by the Coastal Barrier Resources Act (CBRA) (Pub. L. 97-348). "
                details += "CBRA and subsequent amendments (16 U.S.C. § 3501 et seq.) prohibit most Federal funding "
                details += "and financial assistance within System Units, including flood insurance. However, "
                bold = "the existing structure on the property is not within Unit " + cbrs_unit
                bold += " and is therefore not affected by the CBRA restriction on Federal flood insurance."
        elif determination == 5:
            details += "This property is partially within " + system_unit_type + " Unit " + cbrs_unit
            details += " of the CBRS. There is no existing structure on the property. "
            details += "For the portion of the property within " + system_unit_type + " Unit " + cbrs_unit + ", "
            details += opa_info1[0].lower() + opa_info1[1:] + " within OPAs"
            details += opa_info2 if is_opa else su_info[0].lower() + su_info[1:] + " for new construction."
        else:
            details += "A determination has not been made."
        
        if determination in [1, 3]:
            details += " Federal flood insurance through the National Flood Insurance Program is available "
            details += "if the subject building was constructed (or permitted and under construction) before "
            details += "the flood insurance prohibition date, and has not been substantially improved or "
            details += "substantially damaged since. For more information about the restrictions on Federal "
            details += "flood insurance, please refer to the Federal Emergency Management Agency’s (FEMA) "
            details += "regulations in Title 44 Part 71 of the Code of Federal Regulations "
            details += "and FEMA’s Flood Insurance Manual: https://www.fema.gov/flood-insurance-manual."
 
        closing = "We hope this information is helpful. The CBRS maps and additional information can be found "
        closing += " on our website at https://www.fws.gov/cbra/. If you have any additional questions, please contact"
        closing += " Ms. Terri Fish, Program Specialist, at (703) 358-2171.\n"

        signature = "\t\t\t\t\t\tSincerely,\n\n\n\n\n"
        signature += "\t\t\t\t\t\tJonathan Phinney, PhD\n"
        signature += "\t\t\t\t\t\tChief, Branch of Geospatial Mapping and\n"
        signature += "\t\t\t\t\t\tTechnical Support\n"

        cc = ""
        if final_letter_recipient != "":
            cc = "cc:\t" + final_letter_recipient

        # assemble the content sections into a document with proper formatting

        document = Document()
        paragraph_format = document.styles['Normal'].paragraph_format
        paragraph_format.line_spacing = DOC_LINE_SPACING
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1.6)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

        p1 = document.add_paragraph().add_run(referal)
        p1.font.name = DOC_FONT_TYPE
        p1.font.size = DOC_FONT_SMALL

        p2 = document.add_paragraph().add_run(requester_full_address)
        p2.font.name = DOC_FONT_TYPE
        p2.font.size = DOC_FONT_LARGE

        p3 = document.add_paragraph().add_run(salutation)
        p3.font.name = DOC_FONT_TYPE
        p3.font.size = DOC_FONT_LARGE

        p4 = document.add_paragraph().add_run(intro)
        p4.font.name = DOC_FONT_TYPE
        p4.font.size = DOC_FONT_LARGE

        p5 = document.add_paragraph().add_run(property_address)
        p5.font.name = DOC_FONT_TYPE
        p5.font.size = DOC_FONT_LARGE

        p6 = document.add_paragraph().add_run(legal_description)
        p6.font.name = DOC_FONT_TYPE
        p6.font.size = DOC_FONT_LARGE

        # find text between 'https' and '.' and underline/blue font
        details_link = re.search('http(.*).', details)
        p7 = document.add_paragraph()
        if details_link is not None and details_link.group(0):
            details = details.split(details_link.group(0))
            p7_run = p7.add_run(details[0])
            p7_run.font.name = DOC_FONT_TYPE
            p7_run.font.size = DOC_FONT_LARGE
            link = p7.add_run(details_link.group(0))
            link.font.underline = True
            link.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        else:
            p7_run = p7.add_run(details)
            p7_run.font.name = DOC_FONT_TYPE
            p7_run.font.size = DOC_FONT_LARGE

        if bold != "":
            bold = p7.add_run(bold)
            bold.font.bold = True

        closing_link = re.search('http(.*)/.', closing)
        p8 = document.add_paragraph()
        if closing_link is not None and closing_link.group(0):
            closing = closing.split(closing_link.group(0))
            p8_run = p8.add_run(closing[0])
            p8_run.font.name = DOC_FONT_TYPE
            p8_run.font.size = DOC_FONT_LARGE
            link = p8.add_run(closing_link.group(0))
            link.font.underline = True
            link.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            p8.add_run(closing[-1])
        else:
            p8_run = p8.add_run(closing)
            p8_run.font.name = DOC_FONT_TYPE
            p8_run.font.size = DOC_FONT_LARGE

        p9 = document.add_paragraph().add_run(signature)
        p9.font.name = DOC_FONT_TYPE
        p9.font.size = DOC_FONT_LARGE

        p10 = document.add_paragraph().add_run(cc)
        p10.font.name = DOC_FONT_TYPE
        p10.font.size = DOC_FONT_LARGE

        # document.add_paragraph()

        docx_buffer = BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)
        self.document = docx_buffer

        return super(FinalLetterDOCXRenderer, self).render(data, media_type, renderer_context)
